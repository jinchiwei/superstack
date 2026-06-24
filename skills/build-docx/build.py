"""build-docx -- render a markdown file to a Jin-branded DOCX via pandoc.

Usage:
    python build.py --input doc.md --output doc.docx
    python build.py --input doc.md --output doc.docx --double-spaced --sections

Engine: pandoc (system-installed CLI). We pass --reference-doc=reference.docx
which contains all the branded styles. Pandoc applies them to the output.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path


SKILL_DIR = Path(__file__).resolve().parent
REFERENCE_DOCX = SKILL_DIR / "reference.docx"

# Brand byline colors (match build-pdf / build-pptx covers).
_NAME_TURQUOISE = "40E0D0"
_ORG_DEEPPINK = "FF1493"
_DATE_DIM = "888888"


def _parse_frontmatter(md_path: str) -> dict[str, str]:
    """Pull simple key: value pairs from the leading YAML block.

    pandoc only renders title/subtitle/author/date, so name/org/eyebrow are
    otherwise dropped — we re-read them here to inject a branded byline.
    """
    text = Path(md_path).read_text(encoding="utf-8")
    if not text.startswith("---"):
        return {}
    end = text.find("\n---", 3)
    if end == -1:
        return {}
    meta: dict[str, str] = {}
    for line in text[3:end].splitlines():
        m = re.match(r"^([A-Za-z_]+):\s*(.*)$", line)
        if not m:
            continue
        k, v = m.group(1), m.group(2).strip()
        if len(v) >= 2 and v[0] == v[-1] and v[0] in "\"'":
            v = v[1:-1]
        meta[k] = v
    return meta


def _inject_byline(docx_path: str, meta: dict[str, str]) -> None:
    """Insert a branded byline (name turquoise, org deeppink, date dim) right
    after the title/subtitle block. No-op if python-docx is missing or there is
    no name/org to render."""
    name, org, date = meta.get("name"), meta.get("org"), meta.get("date")
    if not (name or org):
        return
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor
        from docx.text.paragraph import Paragraph
    except ImportError:
        print("warning: python-docx not installed; byline injection skipped",
              file=sys.stderr)
        return

    doc = Document(docx_path)
    anchor = None  # last Title or Subtitle paragraph
    for p in doc.paragraphs:
        if p.style is not None and p.style.name in ("Title", "Subtitle"):
            anchor = p
    if anchor is None:
        return

    lines = []
    if name:
        lines.append((name, _NAME_TURQUOISE, True, 14))
    if org:
        lines.append((org, _ORG_DEEPPINK, True, 13))
    if date:
        lines.append((date, _DATE_DIM, False, 11))

    prev_elem, parent = anchor._p, anchor._parent
    for text, color, bold, size in lines:
        new_elem = OxmlElement("w:p")
        prev_elem.addnext(new_elem)
        para = Paragraph(new_elem, parent)
        run = para.add_run(text)
        run.bold = bold
        run.font.name = "Geist"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        prev_elem = new_elem
    doc.save(docx_path)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--input", required=True, help="markdown file to render")
    ap.add_argument("--output", required=True, help="output DOCX path")
    ap.add_argument("--double-spaced", dest="double_spaced", action="store_true",
                    help="2.0 line spacing for journal manuscripts")
    ap.add_argument("--sections", action="store_true",
                    help="number headings (1, 1.1, 1.1.1, ...)")
    ap.add_argument("--toc", action="store_true",
                    help="auto-generated TOC at start (right-click 'update field' in Word)")
    args = ap.parse_args()

    if not shutil.which("pandoc"):
        print("ERROR: pandoc not on PATH. Install via `brew install pandoc`.", file=sys.stderr)
        return 1
    if not REFERENCE_DOCX.is_file():
        print(f"ERROR: {REFERENCE_DOCX} missing. Run make_reference.py to generate.",
              file=sys.stderr)
        return 1

    cmd = [
        "pandoc",
        str(args.input),
        "-o", str(args.output),
        "--reference-doc", str(REFERENCE_DOCX),
        "--from", "markdown+yaml_metadata_block+smart",
        "--to", "docx",
    ]
    if args.toc:
        cmd.append("--toc")
    if args.sections:
        cmd.append("--number-sections")

    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        print(f"pandoc failed:\nSTDERR: {proc.stderr}\nSTDOUT: {proc.stdout}", file=sys.stderr)
        return proc.returncode

    _inject_byline(args.output, _parse_frontmatter(args.input))

    if args.double_spaced:
        _apply_double_spacing(args.output)

    print(f"wrote {args.output}")
    return 0


def _apply_double_spacing(docx_path: str) -> None:
    """Override Normal paragraph spacing to 2.0 in the generated docx."""
    try:
        from docx import Document
    except ImportError:
        print("warning: python-docx not installed; --double-spaced no-op", file=sys.stderr)
        return
    doc = Document(docx_path)
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 2.0
    doc.save(docx_path)


if __name__ == "__main__":
    sys.exit(main())
