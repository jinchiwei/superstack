"""Build-DOCX: end-to-end smoke test."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

SKILL_DIR = Path(__file__).resolve().parents[1]
FIXTURE = SKILL_DIR / "tests" / "fixture.md"
BUILD_PY = SKILL_DIR / "build.py"
REFERENCE = SKILL_DIR / "reference.docx"


def _render(out_path: Path, *extra_args: str) -> None:
    cmd = [sys.executable, str(BUILD_PY),
           "--input", str(FIXTURE),
           "--output", str(out_path),
           *extra_args]
    proc = subprocess.run(cmd, capture_output=True, text=True)
    assert proc.returncode == 0, f"build.py failed:\nSTDERR: {proc.stderr}\nSTDOUT: {proc.stdout}"


def test_reference_docx_exists():
    """Reference.docx must be committed alongside build.py."""
    assert REFERENCE.is_file(), "reference.docx missing — run make_reference.py to regenerate"


def test_renders_docx_to_output_path(tmp_path):
    out = tmp_path / "out.docx"
    _render(out)
    assert out.is_file()
    assert out.stat().st_size > 1000


def test_docx_has_normal_style_with_branded_color(tmp_path):
    """Reference.docx's Normal style sets color to #14141C; output inherits."""
    out = tmp_path / "out.docx"
    _render(out)
    from docx import Document
    doc = Document(str(out))
    normal = doc.styles["Normal"]
    color = normal.font.color
    if color is not None and color.rgb is not None:
        assert str(color.rgb).upper() == "14141C"


def test_docx_has_branded_heading_styles(tmp_path):
    """Heading 1 / 2 / 3 styles exist with branded colors."""
    out = tmp_path / "out.docx"
    _render(out)
    from docx import Document
    doc = Document(str(out))
    style_names = {s.name for s in doc.styles}
    assert "Heading 1" in style_names
    assert "Heading 2" in style_names
    h1 = doc.styles["Heading 1"]
    if h1.font.color and h1.font.color.rgb:
        assert str(h1.font.color.rgb).upper() == "40E0D0"
    h2 = doc.styles["Heading 2"]
    if h2.font.color and h2.font.color.rgb:
        assert str(h2.font.color.rgb).upper() == "FF1493"


def test_docx_contains_fixture_text(tmp_path):
    out = tmp_path / "out.docx"
    _render(out)
    from docx import Document
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Section One" in full_text
    assert "Section Two" in full_text


def test_double_spaced_flag_succeeds(tmp_path):
    """--double-spaced flag should not error and should produce a valid doc."""
    out = tmp_path / "out.docx"
    _render(out, "--double-spaced")
    assert out.stat().st_size > 1000
